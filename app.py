"""
PDF Translator — Streamlit app
================================
Upload a PDF, pick a source and target language, get back a translated
PDF *and* DOCX. Built from the original Colab notebook pipeline:

    PyMuPDF (extract) -> deep-translator/Google Translate (translate)
        -> reportlab + arabic-reshaper + python-bidi (PDF)
        -> python-docx (DOCX)

Run locally:
    streamlit run app.py
"""

import io
import os
import time
import tempfile
from pathlib import Path

import requests
import streamlit as st
import fitz  # PyMuPDF
import arabic_reshaper
from bidi.algorithm import get_display
from deep_translator import GoogleTranslator

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# Language configuration
# ============================================================
# Source language list is intentionally broad: source text is only fed
# into the translator, never rendered, so there's no font/shaping concern.
SOURCE_LANGS = {
    "Auto-detect": "auto", "Hindi": "hi", "English": "en", "Urdu": "ur",
    "Arabic": "ar", "Persian (Farsi)": "fa", "Pashto": "ps", "Sindhi": "sd",
    "Hebrew": "he", "Bengali": "bn", "Tamil": "ta", "Telugu": "te",
    "Marathi": "mr", "Gujarati": "gu", "Punjabi": "pa", "French": "fr",
    "Spanish": "es", "German": "de", "Portuguese": "pt", "Italian": "it",
    "Russian": "ru", "Turkish": "tr", "Chinese (Simplified)": "zh-CN",
    "Japanese": "ja", "Korean": "ko",
}

# Target language list is restricted to languages this app can both
# shape and render correctly in a PDF (no OpenType shaping engine is
# available in reportlab, so complex scripts like Devanagari/CJK are
# left out — see README "Known limitations").
TARGET_LANGS = {
    "Urdu": "ur", "Arabic": "ar", "Persian (Farsi)": "fa", "Pashto": "ps",
    "Sindhi": "sd", "Hebrew": "he", "English": "en", "French": "fr",
    "Spanish": "es", "German": "de", "Portuguese": "pt", "Italian": "it",
    "Russian": "ru", "Turkish": "tr",
}

ARABIC_SCRIPT = {"ur", "ar", "fa", "ps", "sd"}
HEBREW_SCRIPT = {"he"}


def script_for(lang_code: str) -> str:
    if lang_code in ARABIC_SCRIPT:
        return "arabic"
    if lang_code in HEBREW_SCRIPT:
        return "hebrew"
    return "latin"


# Each script gets its own font, downloaded once and cached. Mirrors of
# the same file on two different tags, for resilience if one goes down.
FONT_SOURCES = {
    "arabic": {
        "filename": "NotoNaskhArabic-Regular.ttf",
        "docx_font": "Noto Naskh Arabic",
        "urls": [
            "https://raw.githubusercontent.com/notofonts/notofonts.github.io/main/fonts/NotoNaskhArabic/hinted/ttf/NotoNaskhArabic-Regular.ttf",
            "https://raw.githubusercontent.com/notofonts/notofonts.github.io/noto-monthly-release-2026.05.01/fonts/NotoNaskhArabic/hinted/ttf/NotoNaskhArabic-Regular.ttf",
        ],
    },
    "hebrew": {
        "filename": "NotoSansHebrew-Regular.ttf",
        "docx_font": "Noto Sans Hebrew",
        "urls": [
            "https://raw.githubusercontent.com/notofonts/notofonts.github.io/main/fonts/NotoSansHebrew/hinted/ttf/NotoSansHebrew-Regular.ttf",
            "https://raw.githubusercontent.com/notofonts/notofonts.github.io/noto-monthly-release-2026.05.01/fonts/NotoSansHebrew/hinted/ttf/NotoSansHebrew-Regular.ttf",
        ],
    },
    "latin": {
        "filename": "NotoSans-Regular.ttf",
        "docx_font": None,  # let Word use its default font for Latin/Cyrillic text
        "urls": [
            "https://raw.githubusercontent.com/notofonts/notofonts.github.io/main/fonts/NotoSans/hinted/ttf/NotoSans-Regular.ttf",
            "https://raw.githubusercontent.com/notofonts/notofonts.github.io/noto-monthly-release-2026.05.01/fonts/NotoSans/hinted/ttf/NotoSans-Regular.ttf",
        ],
    },
}

CACHE_DIR = Path(tempfile.gettempdir()) / "pdf_translator_fonts"
CACHE_DIR.mkdir(exist_ok=True)


@st.cache_resource(show_spinner=False)
def get_font_path(script: str) -> str:
    """Download (once, cached) and return the local path to a script's font."""
    # 1. Prefer a font bundled in the repo's fonts/ folder, if present.
    bundled = Path(__file__).parent / "fonts" / FONT_SOURCES[script]["filename"]
    if bundled.exists() and bundled.stat().st_size > 10_000:
        return str(bundled)

    # 2. Otherwise download to a local cache dir (works on Streamlit Cloud).
    dest = CACHE_DIR / FONT_SOURCES[script]["filename"]
    if dest.exists() and dest.stat().st_size > 10_000:
        return str(dest)

    for url in FONT_SOURCES[script]["urls"]:
        try:
            r = requests.get(url, timeout=20)
            if r.ok and len(r.content) > 10_000:
                dest.write_bytes(r.content)
                return str(dest)
        except requests.RequestException:
            continue

    raise RuntimeError(
        f"Could not download the font for script '{script}'. "
        "Check your network connection, or place the .ttf file manually "
        f"in a 'fonts/' folder next to app.py as {FONT_SOURCES[script]['filename']}."
    )


_REGISTERED = set()


def get_pdf_font_name(script: str) -> str:
    """Lazily register a script's font under a stable, script-specific name.

    Using one name per script (rather than a single shared name) avoids a
    race condition if two Streamlit sessions request different scripts at
    the same time in the same server process.
    """
    name = f"PDFFont_{script}"
    if name not in _REGISTERED:
        pdfmetrics.registerFont(TTFont(name, get_font_path(script)))
        _REGISTERED.add(name)
    return name


# ============================================================
# Text extraction
# ============================================================
def extract_pages(pdf_bytes: bytes) -> list[str]:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages = [page.get_text("text").strip() for page in doc]
    doc.close()
    return pages


# ============================================================
# Translation
# ============================================================
MAX_CHUNK = 4500  # stay safely under Google Translate's ~5000-char limit
SLEEP_S = 0.3


def split_into_chunks(text: str, max_len: int) -> list[str]:
    """Split text into <= max_len chunks, preferring paragraph/line boundaries."""
    if len(text) <= max_len:
        return [text]

    chunks = []
    paragraphs = text.split("\n\n")
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 2 <= max_len:
            current = (current + "\n\n" + para).lstrip("\n")
        else:
            if current:
                chunks.append(current)
            if len(para) > max_len:
                lines = para.split("\n")
                sub = ""
                for line in lines:
                    if len(sub) + len(line) + 1 <= max_len:
                        sub = (sub + "\n" + line).lstrip("\n")
                    else:
                        if sub:
                            chunks.append(sub)
                        while len(line) > max_len:
                            chunks.append(line[:max_len])
                            line = line[max_len:]
                        sub = line
                current = sub
            else:
                current = para
    if current:
        chunks.append(current)
    return chunks


def translate_text(text: str, translator: GoogleTranslator) -> str:
    if not text.strip():
        return ""
    chunks = split_into_chunks(text, MAX_CHUNK)
    translated_chunks = []
    for chunk in chunks:
        if not chunk.strip():
            translated_chunks.append("")
            continue
        for attempt in range(4):
            try:
                result = translator.translate(chunk)
                translated_chunks.append(result or "")
                time.sleep(SLEEP_S)
                break
            except Exception:
                time.sleep(3 * (attempt + 1))
        else:
            translated_chunks.append(chunk)  # keep original so page isn't blank
    return "\n\n".join(translated_chunks)


def translate_pages(pages_text: list[str], source: str, target: str, progress_cb) -> list[str]:
    translator = GoogleTranslator(source=source, target=target)
    translated = []
    for i, page in enumerate(pages_text):
        translated.append(translate_text(page, translator))
        progress_cb(i + 1, len(pages_text))
    return translated


# ============================================================
# PDF builder
# ============================================================
PAGE_W, PAGE_H = A4
MARGIN = 20 * mm
FONT_SIZE = 14
LEADING = FONT_SIZE * 1.6
MAX_WIDTH = PAGE_W - 2 * MARGIN


def shape_line(text: str, reshape: bool, rtl: bool) -> str:
    """Reshape (Arabic-script joining) and/or bidi-reorder a line for reportlab."""
    if reshape:
        text = arabic_reshaper.reshape(text)
    if rtl:
        text = get_display(text)
    return text


def wrap_paragraph(text: str, font_name: str, font_size: float, max_width: float, reshape: bool) -> list[str]:
    words = text.split()
    lines, current_words = [], []
    for word in words:
        trial = " ".join(current_words + [word])
        measured = arabic_reshaper.reshape(trial) if reshape else trial
        w = pdfmetrics.stringWidth(measured, font_name, font_size)
        if w <= max_width or not current_words:
            current_words.append(word)
        else:
            lines.append(" ".join(current_words))
            current_words = [word]
    if current_words:
        lines.append(" ".join(current_words))
    return lines


def build_pdf(translated_pages: list[str], target_lang: str) -> bytes:
    script = script_for(target_lang)
    font_name = get_pdf_font_name(script)
    rtl = script in ("arabic", "hebrew")
    reshape = script == "arabic"

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setFont(font_name, FONT_SIZE)
    y = [PAGE_H - MARGIN]
    page_no = [1]

    def footer():
        c.setFont(font_name, 9)
        c.drawCentredString(PAGE_W / 2, MARGIN / 2, str(page_no[0]))
        c.setFont(font_name, FONT_SIZE)

    def new_page():
        footer()
        c.showPage()
        c.setFont(font_name, FONT_SIZE)
        y[0] = PAGE_H - MARGIN
        page_no[0] += 1

    draw_fn = c.drawRightString if rtl else c.drawString
    x_pos = PAGE_W - MARGIN if rtl else MARGIN

    for page_text in translated_pages:
        if not page_text or not page_text.strip():
            continue
        for para in [p for p in page_text.split("\n") if p.strip()]:
            for line in wrap_paragraph(para, font_name, FONT_SIZE, MAX_WIDTH, reshape):
                if y[0] < MARGIN + LEADING:
                    new_page()
                draw_fn(x_pos, y[0], shape_line(line, reshape, rtl))
                y[0] -= LEADING
            y[0] -= LEADING * 0.5

    footer()
    c.save()
    return buf.getvalue()


# ============================================================
# DOCX builder
# ============================================================
def set_paragraph_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def set_run_complex_script_font(run, font_name, size_pt, rtl):
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    if font_name:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:cs"), font_name)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
    if rtl:
        rtl_el = OxmlElement("w:rtl")
        rtl_el.set(qn("w:val"), "1")
        rPr.append(rtl_el)


def add_paragraph(doc, text, font_name, size_pt, rtl):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        set_paragraph_rtl(p)
    run = p.add_run(text)  # original, unshaped text — Word/LibreOffice reshapes itself
    set_run_complex_script_font(run, font_name, size_pt, rtl)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def build_docx(translated_pages: list[str], target_lang: str) -> bytes:
    script = script_for(target_lang)
    rtl = script in ("arabic", "hebrew")
    font_name = FONT_SOURCES[script]["docx_font"]  # None for Latin -> Word default

    doc = Document()
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_p)

    for page_text in translated_pages:
        if not page_text or not page_text.strip():
            continue
        for para in [p for p in page_text.split("\n") if p.strip()]:
            add_paragraph(doc, para, font_name, FONT_SIZE, rtl)
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# Streamlit UI
# ============================================================
st.set_page_config(page_title="PDF Translator", page_icon="📄", layout="centered")

st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,500&family=Inter:wght@400;500;600&display=swap');
    .htu-banner{font-family:Inter,-apple-system,sans-serif;background:#FAF7F2;border:1px solid #E3DACD;
        border-radius:14px;padding:22px 26px;margin:6px 0 16px 0;}
    .htu-banner h1{font-family:Fraunces,Georgia,serif;font-weight:500;font-size:26px;color:#21283B;margin:0 0 4px 0;}
    .htu-banner p{color:#5B6275;margin:0;font-size:14px;}
    .htu-rule{height:3px;width:86px;margin-top:12px;border-radius:2px;
        background:linear-gradient(90deg,#BD8B3F,#21283B,#3F7858);}
    </style>
    <div class="htu-banner">
      <h1>📄 PDF Translator</h1>
      <p>Upload a PDF, translate it, download the result as PDF and DOCX. Powered by Google Translate — free, no API key.</p>
      <div class="htu-rule"></div>
    </div>
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.header("Settings")
    source_name = st.selectbox("Source language", list(SOURCE_LANGS.keys()), index=list(SOURCE_LANGS.keys()).index("Hindi"))
    target_name = st.selectbox("Target language", list(TARGET_LANGS.keys()), index=list(TARGET_LANGS.keys()).index("Urdu"))
    st.caption(
        "Target languages are limited to scripts this app can correctly "
        "shape and render in a PDF (Arabic-script, Hebrew, and Latin-based languages)."
    )

source_code = SOURCE_LANGS[source_name]
target_code = TARGET_LANGS[target_name]

uploaded = st.file_uploader("Choose a PDF file", type=["pdf"])

if "result" not in st.session_state:
    st.session_state.result = None

if uploaded is not None:
    pdf_bytes = uploaded.getvalue()
    base_name = os.path.splitext(uploaded.name)[0]

    if st.button("Translate", type="primary"):
        st.session_state.result = None
        with st.spinner("Extracting text..."):
            pages_text = extract_pages(pdf_bytes)
        total_chars = sum(len(p) for p in pages_text)

        if total_chars == 0:
            st.warning(
                "No text was found in this PDF. It looks like scanned images — "
                "it needs OCR before it can be translated."
            )
        else:
            st.success(f"Extracted {len(pages_text)} page(s), {total_chars:,} characters.")

            progress_bar = st.progress(0.0)
            progress_label = st.empty()

            def progress_cb(done, total):
                progress_bar.progress(done / total)
                progress_label.text(f"{done} / {total} pages translated")

            start = time.time()
            translated_pages = translate_pages(pages_text, source_code, target_code, progress_cb)
            elapsed = time.time() - start
            mins, secs = divmod(int(elapsed), 60)
            st.success(f"Translation complete in {mins}m {secs}s.")

            with st.spinner("Building PDF and DOCX..."):
                pdf_out = build_pdf(translated_pages, target_code)
                docx_out = build_docx(translated_pages, target_code)

            st.session_state.result = {
                "pdf": pdf_out,
                "docx": docx_out,
                "pdf_name": f"Translated_{base_name}.pdf",
                "docx_name": f"Translated_{base_name}.docx",
            }

if st.session_state.result:
    st.subheader("Download")
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "⬇️ Download PDF",
            data=st.session_state.result["pdf"],
            file_name=st.session_state.result["pdf_name"],
            mime="application/pdf",
            use_container_width=True,
        )
    with col2:
        st.download_button(
            "⬇️ Download DOCX",
            data=st.session_state.result["docx"],
            file_name=st.session_state.result["docx_name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

st.caption(
    "Translation is powered by the free, unofficial Google Translate endpoint via deep-translator. "
    "Large files may be slow or rate-limited — this is not an official Google API."
)
